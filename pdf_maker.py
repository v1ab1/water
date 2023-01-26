from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from datetime import datetime

async def pdf_make():
    month = datetime.now().month
    month = month if len(f'{month}') >= 2 else f'0{month}'
    year = f'{datetime.now().year}'[-2:]
    date = f'{datetime.now().day}.{month}.{year}'
    my_file = open("date_pdf.txt", "r")
    dateTxt = my_file.read()
    my_file.close()
    if date == dateTxt:
        return
    my_file = open("date_pdf.txt", "w")
    my_file.write(f'{date}')
    my_file.close()
    def change_heading(heading):
        heading.style.font.size = Pt(24) 
        heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    image_paths = [
        './discounts_parsers/aquamobil/aquamobil_1.png',
        './discounts_parsers/aquamobil/aquamobil_2.png',
        './discounts_parsers/aquamobil/aquamobil_3.png',
        './discounts_parsers/aquamobil/aquamobil_4.jpg',
        "./discounts_parsers/chebarkul/chebarkul_1.jpg",
        "./discounts_parsers/chebarkul/chebarkul_2.jpg",
        "./discounts_parsers/crystal/crystal.png",
        "./discounts_parsers/gorny/gorny_1.jpg",
        "./discounts_parsers/gorny/gorny_2.jpg",
        "./discounts_parsers/gorny/gorny_3.jpg",
        "./discounts_parsers/gorny/gorny_4.jpg",
        "./discounts_parsers/luxe/luxe_0.jpg",
        "./discounts_parsers/luxe/luxe_1.png",
        "./discounts_parsers/luxe/luxe_2.png",
        "./discounts_parsers/luxe/luxe_3.png",
        "./discounts_parsers/luxe/luxe_4.png",
        "./discounts_parsers/luxe/luxe_5.png",
        "./discounts_parsers/luxe/luxe_6.png",
        "./discounts_parsers/luxe/luxe_7.png",
        "./discounts_parsers/luxe/luxe_8.png",
        "./discounts_parsers/luxe/luxe_9.png",
        "./discounts_parsers/luxe/luxe_10.png",
        "./discounts_parsers/luxe/luxe_11.png",
        "./discounts_parsers/luxe/luxe_12.png",
        "./discounts_parsers/luxe/luxe_13.png",
        "./discounts_parsers/luxe/luxe_14.png",
        "./discounts_parsers/luxe/luxe_15.png",
        "./discounts_parsers/luxe/luxe_16.png",
        "./discounts_parsers/niagara/niagara_0.jpg",
        "./discounts_parsers/niagara/niagara_del_1.jpg",
        "./discounts_parsers/niagara/niagara_del_2.jpg",
        "./discounts_parsers/niagara/niagara_del_3.jpg",
        "./discounts_parsers/niagara/niagara_del_4.jpg",
        "./discounts_parsers/niagara/niagara_del_5.jpg",
        "./discounts_parsers/niagara/niagara_del_6.jpg",
        "./discounts_parsers/niagara/niagara_del_7.jpg",
        "./discounts_parsers/niagara/niagara_del_8.jpg",
        "./discounts_parsers/niagara/niagara_store_1.jpg",
        "./discounts_parsers/niagara/niagara_store_2.jpg",
        "./discounts_parsers/niagara/niagara_store_3.jpg",
        "./discounts_parsers/niagara/niagara_store_4.jpg",
        "./discounts_parsers/niagara/niagara_store_5.jpg",
        "./discounts_parsers/niagara/niagara_store_6.jpg",
        "./discounts_parsers/niagara/niagara_store_7.jpg",
        "./discounts_parsers/niagara/niagara_store_8.jpg",
        "./discounts_parsers/zhivaya/zhivaya.png"
    ]

    document = Document()
    title_style = document.styles['Title']
    title_style.font.size = Pt(24)
    for image_path in image_paths:
        if image_path == './discounts_parsers/aquamobil/aquamobil_1.png':
            heading = document.add_heading('Аква-мобиль', level=1)
            change_heading(heading)
        elif image_path == "./discounts_parsers/chebarkul/chebarkul_1.jpg":
            heading = document.add_heading('Чебаркульский исток', level=1)
            change_heading(heading)
        elif image_path == "./discounts_parsers/crystal/crystal.png":
            heading = document.add_heading('Кристальная', level=1)
            change_heading(heading)
        elif image_path == "./discounts_parsers/gorny/gorny_1.jpg":
            heading = document.add_heading('Горный оазис', level=1)
            change_heading(heading)
        elif image_path == "./discounts_parsers/luxe/luxe_0.jpg":
            heading = document.add_heading('Люкс Вода', level=1)
            change_heading(heading)
        elif image_path == "./discounts_parsers/niagara/niagara_0.jpg":
            heading = document.add_heading('Ниагара', level=1)
            change_heading(heading)
        elif image_path == "./discounts_parsers/zhivaya/zhivaya.png":
            heading = document.add_heading('Живая', level=1)
            change_heading(heading)
        document.add_picture(image_path, width=Inches(6))
    document.save("discounts.docx")
    convert("discounts.docx", "discounts.pdf")